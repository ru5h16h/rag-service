from __future__ import annotations

from pathlib import Path

from docx import Document as WordDocument

from src.loaders.base import BaseLoader, Document, build_base_metadata


class DOCXLoader(BaseLoader):
    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def load(self, path: Path) -> list[Document]:
        document = WordDocument(str(path))

        paragraph_text: list[str] = []
        heading_count = 0
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.startswith("Heading"):
                heading_count += 1

            paragraph_text.append(text)

        table_blocks: list[str] = []
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cell_values = [cell.text.strip() for cell in row.cells]
                rows.append("\t".join(cell_values))
            table_blocks.append("\n".join(rows).strip())

        parts = [*paragraph_text, *(block for block in table_blocks if block)]
        text = "\n\n".join(parts).strip()
        if not text:
            return []

        metadata = build_base_metadata(
            path,
            file_format="docx",
            heading_count=heading_count,
            table_count=len(document.tables),
        )
        return [Document(content=text, metadata=metadata)]
